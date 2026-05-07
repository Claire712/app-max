from flask import Flask, render_template_string, request, send_file
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Inches

app = Flask(__name__)

HTML = """
<h2>Rapport d'intervention</h2>

<form method="POST" enctype="multipart/form-data">
  Client : <input name="client"><br><br>
  Technicien : <input name="technicien"><br><br>
  Intervention : <input name="intervention"><br><br>
  Photo : <input type="file" name="photo" accept="image/*"><br><br>

  <button type="submit">Générer rapport</button>
</form>
"""

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        client = request.form["client"]
        technicien = request.form["technicien"]
        intervention = request.form["intervention"]
        photo = request.files.get("photo")

        doc = Document()
        doc.add_heading("Rapport d'intervention", level=1)
        doc.add_paragraph(f"Client : {client}")
        doc.add_paragraph(f"Technicien : {technicien}")
        doc.add_paragraph(f"Intervention : {intervention}")

        if photo and photo.filename:
            filename = secure_filename(photo.filename)
            photo_path = f"uploaded_{filename}"
            photo.save(photo_path)
            doc.add_paragraph("Photo :")
            doc.add_picture(photo_path, width=Inches(4))

        doc.save("rapport.docx")

        return "<h3>Rapport créé ✔️</h3><a href='/download'>Télécharger le rapport</a><br><a href='/'>Retour</a>"

    return render_template_string(HTML)

app.run(debug=True)