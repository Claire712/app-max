from flask import Flask, render_template_string, request
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Inches
import os

app = Flask(__name__)

# Dossier pour stocker les images
UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

HTML = """
<h2>Rapport d'intervention</h2>

<form method="POST" enctype="multipart/form-data">
  Client : <input name="client"><br><br>
  Technicien : <input name="technicien"><br><br>
  Intervention : <input name="intervention"><br><br>

  Photo : <input type="file" name="photo" accept="image/*"><br><br>

  <button type="submit">Générer rapport</button>
</form>
"""

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":

        client = request.form.get("client", "")
        technicien = request.form.get("technicien", "")
        intervention = request.form.get("intervention", "")

        photo = request.files.get("photo")

        # Création du document Word
        doc = Document()
        doc.add_heading("Rapport d'intervention", level=1)

        doc.add_paragraph(f"Client : {client}")
        doc.add_paragraph(f"Technicien : {technicien}")
        doc.add_paragraph(f"Intervention : {intervention}")

        # Gestion photo
        if photo and photo.filename:
            filename = secure_filename(photo.filename)
            photo_path = os.path.join(UPLOAD_FOLDER, filename)
            photo.save(photo_path)

            doc.add_paragraph("Photo :")
            doc.add_picture(photo_path, width=Inches(4))

        # Sauvegarde du rapport
        doc.save("rapport.docx")

        return """
        <h3>Rapport créé ✔️</h3>
        <a href='/'>Retour</a>
        """

    return render_template_string(HTML)


# IMPORTANT POUR RENDER
if __name__ == "__main__":
    app.run(host="0.0.0.0", port=10000)