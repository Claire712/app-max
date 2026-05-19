from flask import Flask, render_template_string, request, send_file
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Cm, Inches, RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

# ==========================================================
# FLASK
# ==========================================================

app = Flask(__name__)

UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
LOGO_PATH = os.path.join(BASE_DIR, "en_tete.png")


# ==========================================================
# HTML
# ==========================================================

HTML = """
<!DOCTYPE html>
<html lang="fr">

<head>
<meta charset="UTF-8">
<title>Rapport intervention</title>

<style>
body{
    font-family: Arial;
    background:#eef2f7;
    padding:40px;
}

.box{
    background:white;
    padding:35px;
    border-radius:18px;
    max-width:700px;
    margin:auto;
    box-shadow:0 0 20px rgba(0,0,0,0.08);
}

h2{
    text-align:center;
    color:#003b73;
    margin-bottom:30px;
}

label{
    font-weight:bold;
    display:block;
    margin-top:15px;
}

input,
select,
textarea{
    width:100%;
    padding:12px;
    margin-top:6px;
    border-radius:8px;
    border:1px solid #ccc;
    box-sizing:border-box;
    font-size:16px;
}

textarea{
    min-height:120px;
}

button{
    background:#005bbb;
    color:white;
    border:none;
    padding:14px 24px;
    border-radius:10px;
    margin-top:25px;
    cursor:pointer;
    font-size:16px;
    width:100%;
}

button:hover{
    background:#004799;
}
</style>
</head>

<body>

<div class="box">

<h2>Rapport d'intervention</h2>

<form method="POST" enctype="multipart/form-data">

<label>Adresse chantier</label>
<input name="chantier">

<label>Syndic</label>
<select name="syndic">
  <option value="Autre" selected>Autre</option>
  <option value="Syndic 1">Syndic 1</option>
  <option value="Syndic 2">Syndic 2</option>
  <option value="Syndic 3">Syndic 3</option>
</select>

<label>Technicien</label>
<input name="technicien">

<label>Date intervention</label>
<input type="date" name="date_intervention">

<label>Objet intervention </label>
<input name="objet">

<label>Compte rendu </label>
<textarea name="intervention"></textarea>

<label>Photos d'intervention</label>

<div id="photo-inputs">
    <input type="file" name="photos" accept="image/*">
</div>

<button type="button"
onclick="addPhotoInput()"
style="margin-top:10px;width:auto;">
Ajouter une photo
</button>

<button type="submit">
Générer rapport
</button>

<script>
function addPhotoInput() {
    const container =
    document.getElementById('photo-inputs');

    const input =
    document.createElement('input');

    input.type = 'file';
    input.name = 'photos';
    input.accept = 'image/*';
    input.style.display = 'block';
    input.style.marginTop = '6px';

    container.appendChild(input);
}
</script>

</form>
</div>
</body>
</html>
"""


# ==========================================================
# ROUTE PRINCIPALE
# ==========================================================

@app.route("/", methods=["GET", "POST"])
def index():

    if request.method == "POST":

        chantier = request.form.get("chantier", "")
        technicien = request.form.get("technicien", "")
        date_intervention = request.form.get(
            "date_intervention", ""
        )
        objet = request.form.get("objet", "")
        intervention = request.form.get(
            "intervention", ""
        )
        syndic = request.form.get("syndic", "Autre")

        photos = request.files.getlist("photos")

        # ==================================================
        # DOCUMENT WORD
        # ==================================================

        doc = Document()

        section = doc.sections[0]
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

        # ==================================================
        # LOGO + SYNDIC
        # ==================================================

        syndic_addresses = {
            "Autre": "Adresse syndic à préciser",
            "Syndic 1": "L'adresse du 1er syndic",
            "Syndic 2": "L'adresse du 2ème syndic",
            "Syndic 3": "Adresse du 3ème syndic",
        }

        syndic_adresse = syndic_addresses.get(
            syndic,
            "Adresse syndic à préciser"
        )

        logo_table = doc.add_table(rows=1, cols=2)
        logo_table.autofit = False

        # logo
        logo_cell = logo_table.cell(0, 0)
        logo_paragraph = logo_cell.paragraphs[0]
        logo_run = logo_paragraph.add_run()

        logo_run.add_picture(
            LOGO_PATH,
            width=Cm(12)
        )

        # syndic
        syndic_cell = logo_table.cell(0, 1)

        syndic_paragraph = syndic_cell.paragraphs[0]
        syndic_paragraph.alignment = (
            WD_PARAGRAPH_ALIGNMENT.RIGHT
        )

        syndic_run = syndic_paragraph.add_run(
            f"Syndic : {syndic}"
        )
        syndic_run.bold = True
        syndic_run.font.size = Pt(17)

        syndic_paragraph2 = (
            syndic_cell.add_paragraph()
        )

        syndic_paragraph2.alignment = (
            WD_PARAGRAPH_ALIGNMENT.RIGHT
        )

        syndic_run2 = syndic_paragraph2.add_run(
            syndic_adresse
        )

        syndic_run2.font.size = Pt(15)

        doc.add_paragraph("")
        doc.add_paragraph("")
        doc.add_paragraph("")

        # ==================================================
        # TITRE
        # ==================================================

        title = doc.add_paragraph()
        title.alignment = (
            WD_PARAGRAPH_ALIGNMENT.CENTER
        )

        t = title.add_run(
            "RAPPORT D'INTERVENTION"
        )

        t.bold = True
        t.font.size = Pt(24)
        t.font.color.rgb = RGBColor(
            0, 70, 140
        )

        doc.add_paragraph("")

        # ==================================================
        # ADRESSE INTERVENTION
        # ==================================================

        adresse_para = doc.add_paragraph()
        adresse_para.alignment = (
            WD_PARAGRAPH_ALIGNMENT.CENTER
        )

        run_adresse = adresse_para.add_run(
            f"Adresse de l'intervention : {chantier}"
        )

        run_adresse.bold = True
        run_adresse.font.size = Pt(13)

        # ==================================================
        # SAUTS
        # ==================================================

        doc.add_paragraph("")
        doc.add_paragraph("")

        # ==================================================
        # PRESTATION
        # ==================================================

        prestation_para = doc.add_paragraph()

        prestation_para.alignment = (
            WD_PARAGRAPH_ALIGNMENT.CENTER
        )

        run_presta_title = (
            prestation_para.add_run(
                "Prestation : "
            )
        )

        run_presta_title.bold = True
        run_presta_title.font.size = Pt(12)

        run_presta = prestation_para.add_run(
            objet
        )

        run_presta.font.size = Pt(12)

        # ==================================================
        # SAUTS
        # ==================================================

        doc.add_paragraph("")
        doc.add_paragraph("")

        # ==================================================
        # DATE
        # ==================================================

        date_para = doc.add_paragraph()

        date_para.alignment = (
            WD_PARAGRAPH_ALIGNMENT.CENTER
        )

        date_title = date_para.add_run(
            "Date intervention : "
        )

        date_title.bold = True
        date_title.font.size = Pt(12)

        date_value = date_para.add_run(
            date_intervention
        )

        date_value.font.size = Pt(12)

        # ==================================================
        # SAUTS
        # ==================================================

        doc.add_paragraph("")
        
        
        # ==================================================
        # TECHNICIEN
        # ==================================================

        doc.add_paragraph("")
        doc.add_paragraph("")

        sign = doc.add_paragraph()

        sign.alignment = (
            WD_PARAGRAPH_ALIGNMENT.CENTER
        )

        sr = sign.add_run(
            f"Technicien : {technicien}"
        )

        sr.bold = True
        sr.font.size = Pt(11)
        doc.add_paragraph("")
        # ==================================================
        # SAUTS
        # ==================================================

        doc.add_paragraph("")
        doc.add_paragraph("")
        
        
        # ==================================================
        # DETAILS INTERVENTION
        # ==================================================

        details = doc.add_paragraph()

        details_title = details.add_run(
             "Détails de l'intervention"
        )

        details_title.bold = True
        details_title.underline = True
        details_title.font.size = Pt(16)
        details_title.font.color.rgb = (
            RGBColor(0, 70, 140)
        )

        texte = doc.add_paragraph(
            intervention
        )

        for run in texte.runs:
            run.font.size = Pt(11)

        # ==================================================
        # SAUTS AVANT PHOTOS
        # ==================================================

        doc.add_paragraph("")
        doc.add_paragraph("")

        # ==================================================
        # PHOTOS
        # ==================================================

        titre_photo = doc.add_paragraph()

        titre_photo.alignment = (
            WD_PARAGRAPH_ALIGNMENT.CENTER
        )

        rr = titre_photo.add_run(
            "PHOTOS D'INTERVENTION"
        )

        rr.bold = True
        rr.underline = True
        rr.font.size = Pt(15)
        rr.font.color.rgb = RGBColor(
            0, 70, 140
        )

        doc.add_paragraph("")

        saved_photos = [
            photo for photo in photos
            if photo and photo.filename
        ]

        current_paragraph = None

        for index, photo in enumerate(
            saved_photos
        ):

            filename = secure_filename(
                photo.filename
            )

            path = os.path.join(
                UPLOAD_FOLDER,
                filename
            )

            photo.save(path)

            if index % 2 == 0:
                current_paragraph = (
                    doc.add_paragraph()
                )

                current_paragraph.alignment = (
                    WD_PARAGRAPH_ALIGNMENT.CENTER
                )

            run = (
                current_paragraph.add_run()
            )

            run.add_picture(
                path,
                width=Inches(2.8)
            )

            run.add_text(" ")

   

        # ==================================================
        # SAVE
        # ==================================================

        filename = (
            "rapport_intervention.docx"
        )

        doc.save(filename)

        return send_file(
            filename,
            as_attachment=True
        )

    return render_template_string(
        HTML
    )


# ==========================================================
# RUN
# ==========================================================

if __name__ == "__main__":

    port = int(os.environ.get("PORT", 10000))

    app.run(
        host="0.0.0.0",
        port=port
    )
